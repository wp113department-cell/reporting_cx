import os
import io
import json
import random
import secrets
import smtplib
import threading
from datetime import datetime, date, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from functools import wraps
from flask import Flask, render_template, request, jsonify, redirect, session, abort
from groq import Groq
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from tinydb import TinyDB, Query
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

# Allow HTTP for local dev; Vercel uses HTTPS automatically
if not os.getenv('VERCEL'):
    os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', os.urandom(24).hex())
app.permanent_session_lifetime = timedelta(days=30)

# On Vercel, filesystem is read-only except /tmp
DB_DIR  = '/tmp/database' if os.getenv('VERCEL') else 'database'
DB_PATH = os.path.join(DB_DIR, 'db.json')
os.makedirs(DB_DIR, exist_ok=True)
db             = TinyDB(DB_PATH)
users_table    = db.table('users')
settings_table = db.table('settings')
records_table  = db.table('records')
tickets_table  = db.table('tickets')

ADMIN_EMAIL       = 'wp113.department@gmail.com'
DRIVE_SCOPES      = ['https://www.googleapis.com/auth/drive.file']
DRIVE_FOLDER_NAME = 'Daily Work Updates'
DB_SYNC_FOLDER    = 'reporting_users'

def _get_creds_path():
    """Return path to credentials.json — from file or GOOGLE_CREDENTIALS_JSON env var."""
    if os.path.exists('credentials.json'):
        return 'credentials.json'
    creds_env = os.getenv('GOOGLE_CREDENTIALS_JSON', '')
    if creds_env:
        tmp = '/tmp/credentials.json'
        with open(tmp, 'w') as f:
            f.write(creds_env)
        return tmp
    return None

def _get_redirect_uri():
    host = request.host
    if 'localhost' in host or '127.0.0.1' in host:
        return f'http://{host}/oauth2callback'
    return f'https://{host}/oauth2callback'

# ─── OTP STORE ───────────────────────────────────────────────────────────────
otp_store = {}  # {email: {'otp': '123456', 'expiry': datetime}}

def generate_otp(email):
    otp = str(random.randint(100000, 999999))
    otp_store[email] = {'otp': otp, 'expiry': datetime.now() + timedelta(minutes=5)}
    print(f"\n{'='*52}")
    print(f"  OTP for {email}:  {otp}")
    print(f"  Expires in 5 minutes")
    print(f"{'='*52}\n")
    return otp

def verify_otp(email, otp):
    if email not in otp_store:
        return False
    stored = otp_store[email]
    if datetime.now() > stored['expiry']:
        otp_store.pop(email, None)
        return False
    if stored['otp'] != otp:
        return False
    otp_store.pop(email, None)
    return True


# ─── DRIVE DB SYNC ───────────────────────────────────────────────────────────
_startup_synced = False

def _get_admin_drive_service_safe():
    try:
        return get_user_drive_service(ADMIN_EMAIL)
    except Exception:
        return None

def _reload_db(content_str):
    global db, users_table, settings_table, records_table, tickets_table
    db.close()
    with open(DB_PATH, 'w', encoding='utf-8') as f:
        f.write(content_str)
    db             = TinyDB(DB_PATH)
    users_table    = db.table('users')
    settings_table = db.table('settings')
    records_table  = db.table('records')
    tickets_table  = db.table('tickets')

def sync_db_to_drive():
    """Push local db.json → Drive/reporting_users/ after every write."""
    try:
        service = _get_admin_drive_service_safe()
        if not service:
            return
        folder_id = get_or_create_folder(service, DB_SYNC_FOLDER)
        with open(DB_PATH, 'rb') as f:
            upload_to_drive(service, folder_id, 'db.json', f.read(), 'application/json')
        print("☁️  DB synced to Drive")
    except Exception as e:
        print(f"⚠️  DB sync to Drive failed: {e}")

def sync_db_from_drive():
    """Pull Drive/reporting_users/db.json → local on startup / login."""
    try:
        service = _get_admin_drive_service_safe()
        if not service:
            return
        q = (f"name='{DB_SYNC_FOLDER}' and "
             f"mimeType='application/vnd.google-apps.folder' and trashed=false")
        folders = service.files().list(q=q, fields='files(id)').execute().get('files', [])
        if not folders:
            return
        q2 = f"name='db.json' and '{folders[0]['id']}' in parents and trashed=false"
        files = service.files().list(q=q2, fields='files(id)').execute().get('files', [])
        if not files:
            return
        raw = service.files().get_media(fileId=files[0]['id']).execute()
        text = raw.decode('utf-8') if isinstance(raw, bytes) else raw
        json.loads(text)        # validate before overwriting
        _reload_db(text)
        print("☁️  DB synced FROM Drive")
    except Exception as e:
        print(f"⚠️  DB sync from Drive failed: {e}")

@app.before_request
def startup_sync():
    global _startup_synced
    if not _startup_synced:
        _startup_synced = True
        threading.Thread(target=sync_db_from_drive, daemon=True).start()


# ─── APPROVAL HELPERS ────────────────────────────────────────────────────────
def get_user_status(email):
    U = Query()
    found = users_table.search(U.email == email)
    return found[0].get('status', 'unknown') if found else 'unknown'

def is_user_approved(email):
    if email == ADMIN_EMAIL:
        return True
    return get_user_status(email) == 'approved'

def _send_raw_email(from_user, from_pass, to_email, subject, html_body):
    brevo_key = os.getenv('BREVO_API_KEY', '')
    if brevo_key:
        import urllib.request as _ureq
        sender_email = from_user or os.getenv('ADMIN_GMAIL_USER', 'noreply@dailyupdate.app')
        payload = json.dumps({
            'sender':      {'name': 'Daily Update App', 'email': sender_email},
            'to':          [{'email': to_email}],
            'subject':     subject,
            'htmlContent': html_body
        }).encode('utf-8')
        req = _ureq.Request('https://api.brevo.com/v3/smtp/email', data=payload,
                             headers={'api-key': brevo_key,
                                      'Content-Type': 'application/json'},
                             method='POST')
        _ureq.urlopen(req, timeout=15)
        return
    # Fallback: Gmail SMTP (local development only)
    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From']    = from_user
    msg['To']      = to_email
    msg.attach(MIMEText(html_body, 'html'))
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as sv:
        sv.login(from_user, from_pass)
        sv.sendmail(from_user, [to_email], msg.as_string())

def send_approval_request(new_user_email, approval_url):
    s         = get_user_settings(ADMIN_EMAIL)
    gmail_user = s.get('gmail_user', '')
    gmail_pass = s.get('gmail_app_password', '')
    html = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;">
  <h2 style="color:#4f46e5;margin-bottom:6px;">🔔 New Access Request</h2>
  <p style="color:#64748b;margin-bottom:20px;">Someone wants access to the <strong>Daily Update Generator</strong>.</p>
  <table style="width:100%;border-collapse:collapse;margin-bottom:24px;">
    <tr>
      <td style="padding:10px 14px;font-weight:bold;background:#f1f5f9;border:1px solid #e2e8f0;width:140px;">User Email</td>
      <td style="padding:10px 14px;border:1px solid #e2e8f0;">{new_user_email}</td>
    </tr>
    <tr>
      <td style="padding:10px 14px;font-weight:bold;background:#f1f5f9;border:1px solid #e2e8f0;">Requested At</td>
      <td style="padding:10px 14px;border:1px solid #e2e8f0;">{datetime.now().strftime('%d/%m/%Y %H:%M')}</td>
    </tr>
  </table>
  <a href="{approval_url}"
     style="display:inline-block;background:#10b981;color:white;padding:14px 36px;border-radius:8px;text-decoration:none;font-weight:bold;font-size:16px;letter-spacing:0.3px;">
    ✅ Approve Access
  </a>
  <p style="color:#94a3b8;font-size:12px;margin-top:20px;">
    Clicking above will approve this user and automatically send them a confirmation email.
  </p>
</div>"""
    if gmail_user and gmail_pass:
        try:
            _send_raw_email(gmail_user, gmail_pass, ADMIN_EMAIL,
                            f'Access Request: {new_user_email}', html)
            print(f"✅ Approval request emailed to admin for {new_user_email}")
            return
        except Exception as e:
            print(f"⚠️  Could not email admin: {e}")
    print(f"\n{'='*60}")
    print(f"  APPROVAL NEEDED for: {new_user_email}")
    print(f"  Approve URL: {approval_url}")
    print(f"{'='*60}\n")

def send_approval_confirmation(user_email):
    s         = get_user_settings(ADMIN_EMAIL)
    gmail_user = s.get('gmail_user', '')
    gmail_pass = s.get('gmail_app_password', '')
    html = """
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;">
  <h2 style="color:#10b981;">🎉 અરજી મંજૂર થઈ ગઈ</h2>
  <p>પ્રિય અરજદાર,</p>
  <p>અભિનંદન! 🎉</p>
  <p>આપ ખરેખર નસીબદાર છો કે માનનીય ભાસ્કર બારોટ સાહેબશ્રીએ આપને આ એપ્લિકેશન વાપરવા લાયક ગણ્યા છે અને આપની અરજી મંજૂર કરી છે.</p>
  <p>હવે આપ આ એપ્લિકેશનનો ઉપયોગ કરી શકો છો.</p>
  <p>સાહેબશ્રીએ આપ પર જે વિશ્વાસ મૂક્યો છે, તેની કદર કરશો અને એપ્લિકેશનનો જવાબદારીપૂર્વક ઉપયોગ કરશો.</p>
  <p>સાહેબશ્રીનો હૃદયપૂર્વક આભાર માનશો. 🙏</p>
  <br>
  <p>શુભેચ્છાઓ,<br><strong>ટીમ</strong></p>
</div>"""
    if gmail_user and gmail_pass:
        try:
            _send_raw_email(gmail_user, gmail_pass, user_email,
                            '🎉 Daily Update App — Access Approved', html)
            print(f"✅ Approval confirmation sent to {user_email}")
            return
        except Exception as e:
            print(f"⚠️  Could not send confirmation to {user_email}: {e}")
    print(f"✅ User {user_email} approved (configure admin Gmail to send confirmation emails)")


# ─── SESSION / AUTH ──────────────────────────────────────────────────────────
def get_current_user():
    return session.get('user_email')

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_email'):
            return redirect('/login')
        return f(*args, **kwargs)
    return decorated

def login_required_api(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_email'):
            return jsonify({'error': 'Not authenticated'}), 401
        return f(*args, **kwargs)
    return decorated


# ─── USER / SETTINGS HELPERS ─────────────────────────────────────────────────
def get_user_settings(email):
    S = Query()
    found = settings_table.search(S.email == email)
    return found[0] if found else {}

def save_user_settings(email, data):
    S = Query()
    data['email'] = email
    if settings_table.search(S.email == email):
        settings_table.update(data, S.email == email)
    else:
        settings_table.insert(data)

def user_setup_done(email):
    s = get_user_settings(email)
    return bool(s.get('setup_done'))

def get_user_groq_client(email):
    s = get_user_settings(email)
    key = s.get('groq_api_key') or os.getenv('GROQ_API_KEY', '')
    return Groq(api_key=key)


def get_user_groq_model(email):
    return 'llama-3.3-70b-versatile'



# ─── DRIVE HELPERS (PER-USER) ─────────────────────────────────────────────────
def get_user_drive_service(email):
    S = Query()
    found = settings_table.search(S.email == email)
    if not found:
        raise Exception("Drive not connected. Go to Settings → Connect Google Drive.")
    token_json = found[0].get('drive_token_json', '')
    if not token_json:
        raise Exception("Drive not connected. Go to Settings → Connect Google Drive.")
    creds = Credentials.from_authorized_user_info(json.loads(token_json), DRIVE_SCOPES)
    if not creds.valid:
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
            settings_table.update({'drive_token_json': creds.to_json()}, S.email == email)
        else:
            raise Exception("Drive session expired. Please reconnect in Settings.")
    return build('drive', 'v3', credentials=creds)

def get_or_create_folder(service, name, parent_id=None):
    q = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        q += f" and '{parent_id}' in parents"
    files = service.files().list(q=q, fields='files(id)').execute().get('files', [])
    if files:
        return files[0]['id']
    meta = {'name': name, 'mimeType': 'application/vnd.google-apps.folder'}
    if parent_id:
        meta['parents'] = [parent_id]
    return service.files().create(body=meta, fields='id').execute()['id']

def upload_to_drive(service, folder_id, filename, content_bytes, mime_type):
    q = f"name='{filename}' and '{folder_id}' in parents and trashed=false"
    existing = service.files().list(q=q, fields='files(id)').execute().get('files', [])
    media = MediaIoBaseUpload(io.BytesIO(content_bytes), mimetype=mime_type, resumable=False)
    if existing:
        file_id = existing[0]['id']
        service.files().update(fileId=file_id, media_body=media).execute()
    else:
        file_id = service.files().create(
            body={'name': filename, 'parents': [folder_id]},
            media_body=media, fields='id'
        ).execute()['id']
    return file_id

def save_to_drive(email, date_key, work_date, form_data, teams_message):
    service    = get_user_drive_service(email)
    root_id    = get_or_create_folder(service, DRIVE_FOLDER_NAME)
    user_folder = get_or_create_folder(service, email, root_id)
    teams_fid  = get_or_create_folder(service, 'Teams', user_folder)
    mail_fid   = get_or_create_folder(service, 'Mail',  user_folder)

    t_id = upload_to_drive(service, teams_fid, f"{date_key}.txt",
                            teams_message.encode('utf-8'), 'text/plain')
    m_id = upload_to_drive(service, mail_fid,  f"{date_key}.docx",
                            _build_docx_bytes(email, work_date, form_data, teams_message),
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return {
        'teams_drive_id':  t_id,
        'mail_drive_id':   m_id,
        'teams_drive_url': f"https://drive.google.com/file/d/{t_id}/view",
        'mail_drive_url':  f"https://drive.google.com/file/d/{m_id}/view",
    }

def fetch_teams_from_drive(email, file_id):
    service = get_user_drive_service(email)
    content = service.files().get_media(fileId=file_id).execute()
    return content.decode('utf-8') if isinstance(content, bytes) else str(content)


# ─── DOCX BUILDER ────────────────────────────────────────────────────────────
def _build_docx_bytes(email, work_date, form_data, teams_message):
    s = get_user_settings(email)
    dept    = s.get('department', 'Production Team')
    manager = s.get('reporting_manager', 'Vipinraj Nair Sir')

    doc   = Document()
    title = doc.add_heading('Daily Work Update', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, value in [
        ('Date', work_date),
        ('Department', dept),
        ('Reporting Manager', manager),
        ('Total Attendance Hours', f"{form_data.get('attendance_hours','N/A')} Hours"),
        ('Total Productive Hours', f"{form_data.get('productive_hours','N/A')} Hours"),
    ]:
        p   = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_paragraph()
    doc.add_heading('Task Summary', level=1)
    tasks = form_data.get('tasks', [])
    if tasks:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, h in enumerate(['Task / Ticket', 'Description of Work Done', 'Productive Hours', 'Status']):
            hdr[i].text = h
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        for task in tasks:
            row = tbl.add_row().cells
            row[0].text = task.get('ticket', 'N/A')
            row[1].text = task.get('description', 'N/A')
            row[2].text = f"{task.get('hours','N/A')} Hours"
            row[3].text = task.get('status', 'N/A')

    meetings = form_data.get('meetings', [])
    doc.add_paragraph()
    doc.add_heading('Meetings', level=1)
    if meetings and any(m.get('name') for m in meetings):
        mt = doc.add_table(rows=1, cols=3)
        mt.style = 'Table Grid'
        mh = mt.rows[0].cells
        for i, h in enumerate(['Meeting Name', 'Duration', 'Purpose']):
            mh[i].text = h
            for para in mh[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        for m in meetings:
            if m.get('name'):
                r = mt.add_row().cells
                r[0].text = m.get('name', 'NA')
                r[1].text = m.get('duration', 'NA')
                r[2].text = m.get('purpose', 'NA')
    else:
        doc.add_paragraph('NA')

    doc.add_paragraph()
    doc.add_heading('Teams Message / Output Summary', level=1)
    for line in teams_message.split('\n'):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── SIMPLE EMAIL HTML ────────────────────────────────────────────────────────
def build_email_html(work_date, form_data, llm_result, user_email):
    s         = get_user_settings(user_email)
    dept      = s.get('department', 'Production Team')
    manager   = s.get('reporting_manager', 'Vipinraj Nair Sir')
    attendance  = form_data.get('attendance_hours', 'N/A')
    productive  = form_data.get('productive_hours', 'N/A')
    task_summaries = llm_result.get('task_summaries', [])
    meetings   = llm_result.get('meetings', [{'name': 'NA', 'duration': 'NA', 'purpose': 'NA'}])

    task_rows = ""
    for ts in task_summaries:
        pts = "".join(f"<li>{p}</li>" for p in ts.get('description_points', []))
        task_rows += f"""<tr>
<td style="padding:8px;border:1px solid #999;font-weight:bold;vertical-align:top;">{ts['ticket']}</td>
<td style="padding:8px;border:1px solid #999;vertical-align:top;"><ul style="margin:0;padding-left:16px;">{pts}</ul></td>
<td style="padding:8px;border:1px solid #999;text-align:center;vertical-align:top;">{ts.get('hours','N/A')}</td>
<td style="padding:8px;border:1px solid #999;text-align:center;vertical-align:top;">{ts.get('status','N/A')}</td>
</tr>"""

    meet_rows = ""
    for m in meetings:
        meet_rows += f"""<tr>
<td style="padding:8px;border:1px solid #999;">{m.get('name','NA')}</td>
<td style="padding:8px;border:1px solid #999;">{m.get('duration','NA')}</td>
<td style="padding:8px;border:1px solid #999;">{m.get('purpose','NA')}</td>
</tr>"""

    summaries = ""
    for ts in task_summaries:
        def li(items): return "".join(f"<li>{x}</li>" for x in items)
        links_html = ""
        if ts.get('links'):
            links_html = "<p><b>Links:</b></p><ul>" + "".join(f"<li>{l}</li>" for l in ts['links']) + "</ul>"
        summaries += f"""
<p><b>{ts['ticket']}</b></p>
{links_html}
<p><b>Commits completed:</b></p><ul>{li(ts.get('commits',['N/A']))}</ul>
<p><b>Modules/screens/features worked on:</b></p><ul>{li(ts.get('modules',['N/A']))}</ul>
<p><b>QA/testing done:</b></p><ul>{li(ts.get('qa_testing',['N/A']))}</ul>
<p><b>Documentation completed:</b></p><ul>{li(ts.get('documentation',['NA']))}</ul>
<p><b>Blockers / Dependencies:</b></p><ul>{li(ts.get('blockers',['None']))}</ul>
<p><b>Tomorrow's Planned Work:</b></p><ul>{li(ts.get('tomorrow',['Continue with the further tasks.']))}</ul>
<hr>"""

    return f"""<div style="font-family:Arial,sans-serif;font-size:14px;color:#000;max-width:800px;">
<p>Date: {work_date}</p>
<p>Department: {dept}</p>
<p>Reporting Manager: {manager}</p>
<p>Total Attendance Hours: {attendance} Hours</p>
<p>Productive Hours: {productive} Hours</p>
<br>
<table border="0" cellpadding="0" cellspacing="0" style="width:100%;border-collapse:collapse;">
<thead><tr style="background:#f0f0f0;">
<th style="padding:8px;border:1px solid #999;text-align:left;">Task / Ticket</th>
<th style="padding:8px;border:1px solid #999;text-align:left;">Description of Work Done</th>
<th style="padding:8px;border:1px solid #999;text-align:center;">Productive Hours Spent</th>
<th style="padding:8px;border:1px solid #999;text-align:center;">Status</th>
</tr></thead>
<tbody>{task_rows}</tbody>
</table>
<br>
<table border="0" cellpadding="0" cellspacing="0" style="width:100%;border-collapse:collapse;">
<thead><tr style="background:#f0f0f0;">
<th style="padding:8px;border:1px solid #999;text-align:left;">Meeting Name</th>
<th style="padding:8px;border:1px solid #999;text-align:left;">Duration</th>
<th style="padding:8px;border:1px solid #999;text-align:left;">Purpose</th>
</tr></thead>
<tbody>{meet_rows}</tbody>
</table>
<br><hr>
<p><b>Output Summary:</b></p>
{summaries}
</div>"""


# ─── LLM (GROQ → GEMINI FALLBACK) ───────────────────────────────────────────
def _parse_json(raw):
    if '```' in raw:
        for part in raw.split('```'):
            s = part.strip()
            if s.startswith('json'):
                raw = s[4:].strip(); break
            elif s.startswith('{'):
                raw = s; break
    return json.loads(raw)

def call_llm(form_data, user_email):
    tasks            = form_data.get('tasks', [])
    attendance_hours = form_data.get('attendance_hours', 'N/A')
    productive_hours = form_data.get('productive_hours', 'N/A')
    work_date        = form_data.get('date', get_working_date())
    meetings         = form_data.get('meetings', [])
    s                = get_user_settings(user_email)
    dept             = s.get('department', 'Production Team')
    manager          = s.get('reporting_manager', 'Vipinraj Nair Sir')

    tasks_text = ""
    for i, task in enumerate(tasks, 1):
        tasks_text += f"""
Task {i}:
- Ticket/Project Name: {task.get('ticket','N/A')}
- What I Did Today: {task.get('description','N/A')}
- Hours Spent: {task.get('hours','N/A')}
- Status: {task.get('status','N/A')}
- Links: {task.get('links','N/A') or 'N/A'}
- Additional Notes: {task.get('notes','N/A') or 'N/A'}
"""
    if meetings and any(m.get('name') for m in meetings):
        meetings_text = "\n".join(f"Name: {m.get('name','NA')} | Duration: {m.get('duration','NA')} | Purpose: {m.get('purpose','NA')}" for m in meetings)
    else:
        meetings_text = "No meetings today."

    prompt = f"""You are a professional AIML developer assistant. Format the daily work update into JSON.

Date: {work_date}
Department: {dept}
Reporting Manager: {manager}
Total Attendance Hours: {attendance_hours} Hours
Total Productive Hours: {productive_hours} Hours

Tasks:
{tasks_text}

Meetings:
{meetings_text}

Return JSON with EXACTLY these fields (no markdown, no extra text):

{{
  "teams_message": "Complete Teams plain-text message",
  "email_subject": "ATTENDANCE: {work_date}",
  "task_summaries": [
    {{
      "ticket": "NAME",
      "description_points": ["point 1","point 2","point 3","point 4"],
      "hours": "X Hours",
      "status": "Status",
      "links": [],
      "commits": ["commit 1","commit 2","commit 3","commit 4","commit 5"],
      "modules": ["module 1","module 2","module 3"],
      "qa_testing": ["test 1","test 2","test 3"],
      "documentation": ["NA"],
      "blockers": ["None"],
      "tomorrow": ["Continue with the further tasks."]
    }}
  ],
  "meetings": [{{"name":"NA","duration":"NA","purpose":"NA"}}]
}}

TEAMS MESSAGE FORMAT:
Work Update

Date: {work_date}

Project/Client Name - [TICKET]

    - [expanded point 1]
    - [expanded point 2]
    - [expanded point 3]
    - [expanded point 4]

[Link: section ONLY if links provided]
Actual Hours:
    - [hours] Hours

[Note: section ONLY if notes provided]
Tomorrow's Task:
    - As per the client's feedback.

RULES:
- Expand user's brief description to 4-6 professional bullet points
- commits: 4-6 past-tense technical actions
- modules: 3-5 specific features worked on
- qa_testing: 3-4 specific tests done
- documentation: ["NA"] unless mentioned
- blockers: ["None"] unless mentioned
- No meetings → [{{"name":"NA","duration":"NA","purpose":"NA"}}]
- Return ONLY valid JSON."""

    client = get_user_groq_client(user_email)
    resp = client.chat.completions.create(
        model=get_user_groq_model(user_email),
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2, max_tokens=4096
    )
    return _parse_json(resp.choices[0].message.content.strip())


# ─── DATE ─────────────────────────────────────────────────────────────────────
def get_working_date():
    today = date.today()
    if today.weekday() == 5:  today -= timedelta(days=1)
    elif today.weekday() == 6: today -= timedelta(days=2)
    return today.strftime('%d/%m/%Y')


# ─── PAGE ROUTES ─────────────────────────────────────────────────────────────
@app.route('/')
@login_required
def index():
    if not user_setup_done(get_current_user()):
        return redirect('/setup')
    return render_template('index.html', user_email=get_current_user())

@app.route('/login')
def login_page():
    if session.get('user_email'):
        return redirect('/')
    return render_template('login.html')

@app.route('/setup')
@login_required
def setup_page():
    s = get_user_settings(get_current_user())
    return render_template('setup.html', settings=s, user_email=get_current_user(),
                           drive_connected=bool(s.get('drive_token_json')))

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')

@app.route('/pending')
def pending_page():
    email = session.get('pending_email', '')
    if not email:
        return redirect('/login')
    if is_user_approved(email):
        session.pop('pending_email', None)
        session.permanent  = True
        session['user_email'] = email
        return redirect('/' if user_setup_done(email) else '/setup')
    return render_template('pending.html', user_email=email)

@app.route('/approve/<token>')
def approve_user(token):
    U = Query()
    found = users_table.search(U.approval_token == token)
    if not found:
        return render_template('approve_result.html', success=False,
                               message='Invalid or expired approval link.')
    user  = found[0]
    email = user['email']
    if user.get('status') == 'approved':
        return render_template('approve_result.html', success=True,
                               message=f'{email} is already approved.')
    users_table.update({'status': 'approved'}, U.approval_token == token)
    sync_db_to_drive()
    send_approval_confirmation(email)
    return render_template('approve_result.html', success=True,
                           message=f'Access approved for {email}. Confirmation sent.')


# ─── AUTH ROUTES ─────────────────────────────────────────────────────────────
@app.route('/api/auth/send-otp', methods=['POST'])
def api_send_otp():
    email = (request.json or {}).get('email', '').strip().lower()
    if not email or '@' not in email:
        return jsonify({'success': False, 'error': 'Invalid email'}), 400

    # Sync from Drive in background (picks up deletions without blocking login)
    threading.Thread(target=sync_db_from_drive, daemon=True).start()

    otp = generate_otp(email)

    # Register user if first time
    U = Query()
    if not users_table.search(U.email == email):
        initial_status = 'approved' if email == ADMIN_EMAIL else 'unknown'
        users_table.insert({'email': email, 'created_at': datetime.now().isoformat(),
                            'status': initial_status})
        sync_db_to_drive()  # push new user to Drive immediately

    # Send OTP via email — use user's own Gmail if configured, else fall back to admin's Gmail
    s          = get_user_settings(email)
    gmail_user = s.get('gmail_user', '') or ''
    gmail_pass = s.get('gmail_app_password', '') or ''
    if not (gmail_user and gmail_pass):
        admin_s    = get_user_settings(ADMIN_EMAIL)
        gmail_user = admin_s.get('gmail_user', '')
        gmail_pass = admin_s.get('gmail_app_password', '')
    # Final fallback: env vars (needed on fresh Render/Vercel where DB is empty)
    if not (gmail_user and gmail_pass):
        gmail_user = os.getenv('ADMIN_GMAIL_USER', '')
        gmail_pass = os.getenv('ADMIN_GMAIL_PASS', '')

    sent_email = False
    otp_html = f"""
<div style="font-family:Arial,sans-serif;max-width:480px;margin:0 auto;padding:24px;">
  <h2 style="color:#4f46e5;margin-bottom:4px;">🔐 Your Login OTP</h2>
  <p style="color:#64748b;margin-bottom:20px;">Daily Update Generator</p>
  <div style="background:#f1f5f9;border-radius:12px;padding:24px;text-align:center;margin-bottom:20px;">
    <div style="font-size:36px;font-weight:900;letter-spacing:10px;color:#1e293b;">{otp}</div>
  </div>
  <p style="color:#64748b;font-size:13px;">This OTP expires in <strong>5 minutes</strong>. Do not share it with anyone.</p>
</div>"""
    try:
        _send_raw_email(gmail_user, gmail_pass, email,
                        'Your OTP — Daily Update App', otp_html)
        sent_email = True
    except Exception as e:
        print(f"⚠️  OTP email failed: {e}")

    msg = 'OTP sent to your email!' if sent_email else 'OTP printed in server terminal (check the terminal window).'
    return jsonify({'success': True, 'message': msg})

@app.route('/api/auth/verify-otp', methods=['POST'])
def api_verify_otp():
    data  = request.json or {}
    email = data.get('email', '').strip().lower()
    otp   = data.get('otp', '').strip()

    if not verify_otp(email, otp):
        return jsonify({'success': False, 'error': 'Invalid or expired OTP'}), 400

    # Approval gate — admin always passes
    if not is_user_approved(email):
        status = get_user_status(email)
        if status != 'pending':
            # First access request — generate token and email admin
            token = secrets.token_urlsafe(32)
            U2 = Query()
            users_table.update({'status': 'pending', 'approval_token': token}, U2.email == email)
            sync_db_to_drive()
            approval_url = f"{request.host_url.rstrip('/')}/approve/{token}"
            send_approval_request(email, approval_url)
        session['pending_email'] = email
        return jsonify({'success': True, 'status': 'pending', 'redirect': '/pending'})

    session.permanent = True
    session['user_email'] = email
    setup_done = user_setup_done(email)
    return jsonify({'success': True, 'setup_done': setup_done,
                    'redirect': '/' if setup_done else '/setup'})


# ─── SETTINGS ROUTES ─────────────────────────────────────────────────────────
@app.route('/api/settings/save', methods=['POST'])
@login_required_api
def api_settings_save():
    email = get_current_user()
    data  = request.json or {}
    data['setup_done'] = True
    save_user_settings(email, data)
    sync_db_to_drive()
    return jsonify({'success': True, 'message': 'Settings saved!'})

@app.route('/api/settings/get')
@login_required_api
def api_settings_get():
    email = get_current_user()
    s = get_user_settings(email)
    safe = {k: v for k, v in s.items() if k not in ('gmail_app_password', 'drive_token_json', 'email')}
    has_drive = bool(s.get('drive_token_json'))
    return jsonify({'success': True, 'settings': safe, 'drive_connected': has_drive})


# ─── DRIVE AUTH ──────────────────────────────────────────────────────────────
@app.route('/api/drive/status')
@login_required_api
def api_drive_status():
    email = get_current_user()
    try:
        get_user_drive_service(email)
        return jsonify({'connected': True})
    except Exception as e:
        return jsonify({'connected': False, 'error': str(e)})

@app.route('/api/validate/groq', methods=['POST'])
@login_required_api
def api_validate_groq():
    data  = request.json or {}
    key   = data.get('api_key', '').strip()
    model = data.get('model', 'llama-3.3-70b-versatile').strip() or 'llama-3.3-70b-versatile'
    if not key:
        return jsonify({'valid': False, 'error': 'No API key provided'})
    try:
        client = Groq(api_key=key)
        client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": "Reply with just: OK"}],
            max_tokens=5, temperature=0
        )
        return jsonify({'valid': True, 'message': f'Groq key is valid! Model "{model}" works.'})
    except Exception as e:
        return jsonify({'valid': False, 'error': str(e)})


@app.route('/api/validate/gmail', methods=['POST'])
@login_required_api
def api_validate_gmail():
    data       = request.json or {}
    gmail_user = data.get('gmail_user', '').strip()
    gmail_pass = data.get('gmail_app_password', '').strip()
    if not gmail_user or not gmail_pass:
        return jsonify({'valid': False, 'error': 'Email and App Password are required'})
    # On cloud servers SMTP is blocked — credentials are saved and Brevo handles delivery
    if os.getenv('BREVO_API_KEY'):
        return jsonify({'valid': True, 'message': 'Credentials saved! Emails are sent via Brevo (SMTP blocked on cloud).'})
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465, timeout=10) as server:
            server.login(gmail_user, gmail_pass)
        return jsonify({'valid': True, 'message': 'Gmail connection successful!'})
    except smtplib.SMTPAuthenticationError:
        return jsonify({'valid': False, 'error': 'Authentication failed. Check your Gmail address and App Password.'})
    except Exception as e:
        return jsonify({'valid': False, 'error': str(e)})


@app.route('/api/drive/start-auth', methods=['POST'])
@login_required_api
def api_drive_start_auth():
    email = get_current_user()
    creds_path = _get_creds_path()
    if not creds_path:
        return jsonify({'success': False, 'error': 'credentials.json not found. Set GOOGLE_CREDENTIALS_JSON env var.'}), 400
    try:
        redirect_uri = _get_redirect_uri()
        flow = Flow.from_client_secrets_file(creds_path, DRIVE_SCOPES,
                                                          redirect_uri=redirect_uri)
        auth_url, state = flow.authorization_url(access_type='offline', prompt='consent')
        session['oauth_mode']        = 'drive'
        session['drive_auth_email']  = email
        session['drive_auth_state']  = state
        return jsonify({'success': True, 'auth_url': auth_url})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/oauth2callback')
def oauth2callback():
    email = session.get('drive_auth_email')
    if not email:
        return redirect('/setup?error=auth_expired')
    creds_path = _get_creds_path()
    if not creds_path:
        return redirect('/setup?error=no_credentials')
    try:
        redirect_uri = _get_redirect_uri()
        flow = Flow.from_client_secrets_file(
            creds_path, DRIVE_SCOPES, redirect_uri=redirect_uri)
        flow.fetch_token(authorization_response=request.url)
        creds = flow.credentials
        S = Query()
        if settings_table.search(S.email == email):
            settings_table.update({'drive_token_json': creds.to_json()}, S.email == email)
        else:
            settings_table.insert({'email': email, 'drive_token_json': creds.to_json()})
        sync_db_to_drive()
        session.pop('drive_auth_email', None)
        session.pop('drive_auth_state', None)
        session.pop('oauth_mode', None)
        return redirect('/setup?drive=connected')
    except Exception as e:
        print(f"Drive auth error: {e}")
        return redirect('/setup?error=drive_failed')


# ─── MAIN APP API ─────────────────────────────────────────────────────────────
@app.route('/api/date')
def api_date():
    return jsonify({'date': get_working_date()})

@app.route('/api/tickets')
@login_required_api
def api_tickets():
    email = get_current_user()
    T = Query()
    history_names = set(t['name'] for t in tickets_table.search(T.user == email))
    s = get_user_settings(email)
    saved = [p.strip() for p in s.get('project_names', '').splitlines() if p.strip()]
    all_names = sorted(set(list(history_names) + saved))
    return jsonify({'tickets': all_names, 'pinned': saved})

@app.route('/api/generate', methods=['POST'])
@login_required_api
def api_generate():
    try:
        email     = get_current_user()
        form_data = request.json
        if not form_data.get('tasks'):
            return jsonify({'success': False, 'error': 'No tasks provided'}), 400

        llm_result = call_llm(form_data, email)
        work_date  = form_data.get('date', get_working_date())
        email_html = build_email_html(work_date, form_data, llm_result, email)

        T = Query()
        for task in form_data.get('tasks', []):
            name = task.get('ticket', '').strip()
            if name and not tickets_table.search((T.name == name) & (T.user == email)):
                tickets_table.insert({'name': name, 'user': email})

        return jsonify({
            'success':       True,
            'teams_message': llm_result.get('teams_message', ''),
            'email_html':    email_html,
            'email_subject': llm_result.get('email_subject', f"ATTENDANCE: {work_date}")
        })
    except json.JSONDecodeError as e:
        return jsonify({'success': False, 'error': f'LLM returned invalid JSON. Try again. ({e})'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/save', methods=['POST'])
@login_required_api
def api_save():
    try:
        email        = get_current_user()
        data         = request.json
        form_data    = data.get('form_data', {})
        teams_msg    = data.get('teams_message', '')
        work_date    = form_data.get('date', get_working_date())
        date_key     = work_date.replace('/', '-')

        drive_result = save_to_drive(email, date_key, work_date, form_data, teams_msg)

        R = Query()
        record = {
            'user':            email,
            'date':            date_key,
            'display_date':    work_date,
            'teams_drive_id':  drive_result['teams_drive_id'],
            'mail_drive_id':   drive_result['mail_drive_id'],
            'teams_drive_url': drive_result['teams_drive_url'],
            'mail_drive_url':  drive_result['mail_drive_url'],
            'task_count':      len(form_data.get('tasks', [])),
            'tickets':         [t.get('ticket','') for t in form_data.get('tasks', [])],
            'saved_at':        datetime.now().isoformat()
        }
        if records_table.search((R.date == date_key) & (R.user == email)):
            records_table.update(record, (R.date == date_key) & (R.user == email))
        else:
            records_table.insert(record)
        sync_db_to_drive()

        return jsonify({'success': True, 'message': 'Saved to Google Drive!', 'drive': drive_result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/send_email', methods=['POST'])
@login_required_api
def api_send_email():
    try:
        email = get_current_user()
        s     = get_user_settings(email)
        gmail_user = s.get('gmail_user', '')
        gmail_pass = s.get('gmail_app_password', '')
        to_email   = s.get('to_email', '')
        cc_raw     = s.get('cc_emails', '')
        cc_emails  = [x.strip() for x in cc_raw.split(',') if x.strip()] if isinstance(cc_raw, str) else cc_raw

        if not gmail_user or not gmail_pass:
            return jsonify({'success': False, 'error': 'Gmail not configured. Go to Settings.'}), 400
        if not to_email:
            return jsonify({'success': False, 'error': 'To email not configured. Go to Settings.'}), 400

        data          = request.json
        email_html    = data.get('email_html', '')
        email_subject = data.get('email_subject', 'Daily Work Update')

        full_html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"></head>
<body style="margin:0;padding:20px;font-family:Arial,sans-serif;">{email_html}</body></html>"""

        all_recipients = [to_email] + cc_emails

        brevo_key = os.getenv('BREVO_API_KEY', '')
        if brevo_key:
            import urllib.request as _ureq
            for recipient in all_recipients:
                payload = json.dumps({
                    'sender':      {'name': 'Daily Update App', 'email': gmail_user},
                    'to':          [{'email': recipient}],
                    'subject':     email_subject,
                    'htmlContent': full_html
                }).encode('utf-8')
                req = _ureq.Request('https://api.brevo.com/v3/smtp/email', data=payload,
                                     headers={'api-key': brevo_key,
                                              'Content-Type': 'application/json'},
                                     method='POST')
                _ureq.urlopen(req, timeout=15)
        else:
            msg = MIMEMultipart('alternative')
            msg['Subject'] = email_subject
            msg['From']    = gmail_user
            msg['To']      = to_email
            if cc_emails:
                msg['Cc'] = ', '.join(cc_emails)
            msg.attach(MIMEText(full_html, 'html'))
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                server.login(gmail_user, gmail_pass)
                server.sendmail(gmail_user, all_recipients, msg.as_string())

        return jsonify({'success': True, 'message': 'Email sent!'})
    except smtplib.SMTPAuthenticationError:
        return jsonify({'success': False, 'error': 'Gmail auth failed. Check your App Password in Settings.'}), 401
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history')
@login_required_api
def api_history():
    email = get_current_user()
    R     = Query()
    records = records_table.search(R.user == email)
    records.sort(key=lambda x: x.get('date', ''), reverse=True)
    return jsonify({'records': records})

@app.route('/api/history/<date_key>')
@login_required_api
def api_history_detail(date_key):
    email = get_current_user()
    R     = Query()
    found = records_table.search((R.date == date_key) & (R.user == email))
    if not found:
        return jsonify({'error': 'Not found'}), 404
    record = found[0]
    teams_content = ''
    file_id = record.get('teams_drive_id', '')
    if file_id:
        try:
            teams_content = fetch_teams_from_drive(email, file_id)
        except Exception as e:
            teams_content = f'[Could not fetch from Drive: {e}]'
    return jsonify({'record': record, 'teams_content': teams_content})

@app.route('/api/download/<file_type>/<date_key>')
@login_required_api
def api_download(file_type, date_key):
    email = get_current_user()
    R     = Query()
    found = records_table.search((R.date == date_key) & (R.user == email))
    if not found:
        abort(404)
    record = found[0]
    url = record.get('teams_drive_url') if file_type == 'teams' else record.get('mail_drive_url')
    if not url:
        abort(404)
    return redirect(url)


if __name__ == '__main__':
    print("\n" + "="*50)
    print("  Daily Work Update Generator")
    print("  http://localhost:5000")
    print("="*50 + "\n")
    debug = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
    port = int(os.getenv('PORT', 5000))
    app.run(debug=debug, host='0.0.0.0', port=port)
